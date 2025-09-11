# STEP 1: Load .env file
from dotenv import load_dotenv

load_dotenv()

# STEP 2: Import all libraries
import os
import json
import io
import csv
from googletrans import Translator
from docx import Document
from fpdf import FPDF
from fastapi import FastAPI, File, Form, UploadFile, HTTPException, Request, Body, Depends
app = FastAPI()
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
import chromadb
import requests
import psycopg2
from psycopg2 import pool
from psycopg2.extras import RealDictCursor
from datetime import datetime
from typing import Optional
from google.oauth2 import id_token
from google.auth.transport import requests as google_requests
from pathlib import Path
import re

# =========================
# CONFIG
# =========================
OLLAMA_API_URL = "http://localhost:11434/api/chat"
MODEL_NAME = "llama3.2:3b"
GOOGLE_CLIENT_ID = "226312071852-bpt8lnl56pkh0uf544bu3ufk604fms9r.apps.googleusercontent.com"

from urllib.parse import urlparse, parse_qsl, urlencode, urlunparse

def _sanitize_dsn(dsn: str) -> str:
    """Drop query params psycopg2 doesn't understand (e.g., pgbouncer=true)."""
    try:
        parsed = urlparse(dsn)
        q = dict(parse_qsl(parsed.query, keep_blank_values=True))
        # Remove Node/Prisma-specific flag that psycopg2 can't parse
        q.pop("pgbouncer", None)
        new_query = urlencode(q)
        return urlunparse(parsed._replace(query=new_query))
    except Exception:
        return dsn

# Gather DB candidates (in a resilient order) and try until one works.
env_candidates = []
if os.getenv("PY_DATABASE_URL"):
    env_candidates.append(("PY_DATABASE_URL", os.getenv("PY_DATABASE_URL")))
if os.getenv("DATABASE_URL"):
    env_candidates.append(("DATABASE_URL", os.getenv("DATABASE_URL")))
if os.getenv("DIRECT_URL"):
    env_candidates.append(("DIRECT_URL", os.getenv("DIRECT_URL")))

if not env_candidates:
    raise EnvironmentError("No database URL found. Set PY_DATABASE_URL or DATABASE_URL or DIRECT_URL")

last_err = None
db_pool = None
for name, raw in env_candidates:
    dsn = _sanitize_dsn(raw)
    try:
        # Quick connectivity probe to catch DNS issues early
        test_conn = psycopg2.connect(dsn)
        test_conn.close()
        db_pool = pool.SimpleConnectionPool(1, 20, dsn=dsn, cursor_factory=RealDictCursor)
        print(f"[backend] Connected to DB via {name}")
        break
    except Exception as e:
        last_err = e
        print(f"[backend] DB connect failed using {name}: {e}")

if db_pool is None:
    raise last_err or EnvironmentError("Failed to initialize database pool")

chroma_client = chromadb.PersistentClient(path="chroma_db")
collection = chroma_client.get_or_create_collection(name="cybersecurity")

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =========================
# MODELS
# =========================
class PromptRequest(BaseModel):
    prompt: str
    user_id: str
    session_id: int = None

class SessionCreateRequest(BaseModel):
    user_id: str
    title: str | None = None

class GoogleToken(BaseModel):
    token: str

class FeedbackBody(BaseModel):
    user_id: Optional[str] = None
    rating: Optional[int] = None
    category: Optional[str] = None
    message: str
    contact_email: Optional[str] = None

# =========================
# HELPERS
# =========================
def get_db_connection():
    """Dependency to get a DB connection from the pool."""
    conn = None
    try:
        conn = db_pool.getconn()
        yield conn
    finally:
        if conn:
            db_pool.putconn(conn)

def retrieve_context(query: str, n_results: int = 3) -> str:
    try:
        results = collection.query(query_texts=[query], n_results=n_results)
        if results and "documents" in results and results["documents"]:
            retrieved_docs = [doc for docs in results["documents"] for doc in docs]
            context = "\n\n".join(retrieved_docs)
            return context.strip()
    except Exception as e:
        print(f"❌ Error querying ChromaDB: {e}")
    return "No relevant context found."

def _is_greeting(text: str) -> bool:
    """Detect simple greeting-only inputs like 'hi', 'hello', 'good morning'."""
    try:
        t = (text or "").strip().lower()
        # keep only letters and spaces
        t = re.sub(r"[^a-z\s]", "", t)
        words = [w for w in t.split() if w]
        if not words:
            return False
        # If the input is short and all words are in greeting set, treat as greeting
        if len(words) > 4:
            return False
        greeting_words = {
            "hi", "hello", "hey", "yo", "hiya", "hola", "sup",
            "good", "morning", "afternoon", "evening", "there"
        }
        return all(w in greeting_words for w in words)
    except Exception:
        return False

def _style_instructions(style: str | None) -> str:
    try:
        s = (style or "long").strip().lower()
    except Exception:
        s = "long"
    if s in {"summary", "summarize"}:
        return (
            "STYLE: Provide a brief overview only. One to three short sentences after the title. "
            "If listing, include at most three hyphen bullets, one per line. No numbered section."
        )
    if s in {"short", "short answer"}:
        return (
            "STYLE: Respond in one or two short sentences after the title. "
            "Avoid lists and examples unless strictly necessary."
        )
    if s in {"main", "main points", "only main point", "only main points"}:
        return (
            "STYLE: Return only the essential bullet points. Use 4-7 hyphen bullets, one per line. "
            "No extra prose before or after."
        )
    return (
        "STYLE: Provide a detailed answer. Include Essential Steps (hyphen bullets) and Advanced Measures (numbered lines). "
        "Add commands/code if helpful, and a short note and references when relevant."
    )


def call_llm(prompt: str, history: list[dict] | None = None, style: str | None = None) -> str:
    """Gets a single, complete response from the LLM with translation."""
    english_prompt, source_lang = _translate(prompt, 'en')

    # Short-circuit for simple greetings
    if _is_greeting(english_prompt):
        english_answer = "Hello! How can I help you with cybersecurity today?"
        final_answer, _ = _translate(english_answer, source_lang)
        return final_answer

    context = retrieve_context(english_prompt)
    system_prompt = (
        "You are a professional cybersecurity assistant. "
        "Write in plain text with minimal Markdown ONLY for code blocks and blockquotes. Do NOT use heading markers (# or ##). Do NOT use asterisks (*) for bold/italics. Keep sentences short and place each sentence on its own line. Leave a blank line between sections.\n\n"
        "Start with a single TITLE line that states the topic . No markup.\n"
        "Immediately after the title, write a one- or two-sentence overview WITHOUT any label like 'Summary'. Each sentence on its own line.\n\n"
        "Then use these sections and styles:\n\n"
        "Essential Steps\n"
        "- Hyphen bullets. 3–6 items. One sentence per bullet. Put each bullet on its own line. Do not join bullets on the same line.\n\n"
        "Advanced Measures\n"
        "1. Step title on this line.\n\n"
        "2. Next step title on this line.\n\n"
        "3. Next step title on this line.\n\n"
        "Use the exact numbering style with a period (e.g., 1.) and put each numbered item on its own line. Do not join multiple numbers on one line. Add a blank line after each numbered item. Sub-points under a numbered step may use hyphen bullets.\n\n"
        "```bash\n<commands or code here>\n```\n\n"
        "> Important notes or warnings should be provided as blockquote lines beginning with '>'.\n\n"
        "References (optional)\n"
        "- Links or document names.\n\n"
        f"{_style_instructions(style)}\n"
        "Always ground answers in the relevant context below when helpful. Prefer concrete actions over theory."
        "if the question if unrelated to cybersecurity, politely inform the user that you are specialized in cybersecurity topics and cannot assist with their query."
        f"\n\n--- Relevant Context ---\n{context}"
    )
    messages = [
        {"role": "system", "content": system_prompt},
        *(history or []),
        {"role": "user", "content": english_prompt}
    ]
    payload = {"model": MODEL_NAME, "messages": messages, "stream": False}

    try:
        response = requests.post(OLLAMA_API_URL, json=payload, timeout=120)
        response.raise_for_status()
        data = response.json()
        english_answer = data.get("message", {}).get("content", "No response from model.")

        final_answer, _ = _translate(english_answer, source_lang)
        return final_answer
    except Exception as e:
        print(f"Error calling local Ollama LLM: {e}")
        return "Sorry, I couldn't process your request."

def stream_llm_response(prompt: str, history: list[dict] | None = None, style: str | None = None):
    """A generator function that streams the response from the LLM with translation."""
    english_prompt, source_lang = _translate(prompt, 'en')

    # Short-circuit for simple greetings
    if _is_greeting(english_prompt):
        english_answer = "Hello! How can I help you with cybersecurity today?"
        final_answer, _ = _translate(english_answer, source_lang)
        for word in final_answer.split():
            yield word + " "
        return

    context = retrieve_context(english_prompt)
    system_prompt = (
        "You are a professional cybersecurity assistant. "
        "Write in plain text with minimal Markdown ONLY for code blocks and blockquotes. Do NOT use heading markers (# or ##). Do NOT use asterisks (*) for bold/italics. Keep sentences short and place each sentence on its own line. Leave a blank line between sections.\n\n"
        "Start with a single TITLE line that states the topic . No markup.\n"
        "Immediately after the title, write a one- or two-sentence overview WITHOUT any label like 'Summary'. Each sentence on its own line.\n\n"
        "Then use these sections and styles:\n\n"
        "Essential Steps\n"
        "- Hyphen bullets. 3–6 items. One sentence per bullet. Put each bullet on its own line. Do not join bullets on the same line.\n\n"
        "Advanced Measures\n"
        "1. Step title on this line.\n\n"
        "2. Next step title on this line.\n\n"
        "3. Next step title on this line.\n\n"
        "Use the exact numbering style with a period (e.g., 1.) and put each numbered item on its own line. Do not join multiple numbers on one line. Add a blank line after each numbered item. Sub-points under a numbered step may use hyphen bullets.\n\n"
        "```bash\n<commands or code here>\n```\n\n"
        "> Important notes or warnings should be provided as blockquote lines beginning with '>'.\n\n"
        "References (optional)\n"
        "- Links or document names.\n\n"
        f"{_style_instructions(style)}\n"
        "Always ground answers in the relevant context below when helpful. Prefer concrete actions over theory."
        f"\n\n--- Relevant Context ---\n{context}"
    )
    messages = [
        {"role": "system", "content": system_prompt},
        *(history or []),
        {"role": "user", "content": english_prompt}
    ]
    payload = {"model": MODEL_NAME, "messages": messages, "stream": True}

    try:
        response = requests.post(OLLAMA_API_URL, json=payload, stream=True)
        response.raise_for_status()

        english_answer = ""
        for chunk in response.iter_lines():
            if chunk:
                data = json.loads(chunk)
                english_answer_chunk = data.get("message", {}).get("content", "")
                english_answer += english_answer_chunk

        final_answer, _ = _translate(english_answer, source_lang)

        for word in final_answer.split():
            yield word + " "

    except Exception as e:
        print(f"Error during streaming: {e}")
        yield "Sorry, an error occurred during streaming."


# --- DOWNLOAD HELPERS ---
def _create_docx(messages):
    document = Document()
    document.add_heading('Chat History', 0)
    for msg in messages:
        p = document.add_paragraph()
        p.add_run(f'{msg["role"].capitalize()}: ').bold = True
        p.add_run(msg["content"])
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def _create_pdf(messages):
    try:
        pdf = FPDF()
        pdf.add_page()

        font_path = Path(__file__).parent / "fonts" / "DejaVuSans.ttf"
        pdf.add_font("DejaVu", "", str(font_path))
        pdf.set_font("DejaVu", size=12)

        pdf.cell(0, 10, txt="Chat History", ln=True, align='C')
        pdf.ln(5)

        for msg in messages:
            pdf.set_font("DejaVu", size=10)
            pdf.write(8, f'{msg["role"].capitalize()}: ')
            pdf.set_font("DejaVu", size=10)
            pdf.write(8, msg["content"])
            pdf.ln(12)

        return io.BytesIO(pdf.output())
    except Exception as e:
        print(f"🔴 FAILED TO CREATE PDF: {e}")
        raise e

def _create_csv(messages):
    file_stream = io.StringIO()
    writer = csv.writer(file_stream)
    writer.writerow(['role', 'content'])
    for msg in messages:
        writer.writerow([msg['role'], msg['content']])
    file_stream.seek(0)
    return io.BytesIO(file_stream.read().encode('utf-8'))

def _translate(text: str, dest_lang: str):
    """Translates text to a destination language and detects the source."""
    try:
        translator = Translator()
        detected_lang = translator.detect(text).lang
        if detected_lang == dest_lang:
            return text, detected_lang

        translated = translator.translate(text, dest=dest_lang)
        return translated.text, detected_lang
    except Exception as e:
        print(f"Error during translation: {e}")
        return text, 'en'

# =========================
# ROUTES
# =========================

@app.post("/chat/session")
def create_chat_session(user_id: str = Form(...), title: str = Form("New Chat"), conn=Depends(get_db_connection)):
    """Create a new chat session for a user."""
    with conn.cursor() as cursor:
        cursor.execute(
            "INSERT INTO chat_sessions (user_id, title) VALUES (%s, %s) RETURNING id, created_at",
            (user_id, title or "New Chat"),
        )
        row = cursor.fetchone()
        conn.commit()
        return {"session_id": row.get("id"), "title": title or "New Chat", "created_at": row.get("created_at")}

@app.get("/chat/sessions/{user_id}")
def list_chat_sessions(user_id: str, conn=Depends(get_db_connection)):
    """List chat sessions for a user."""
    with conn.cursor() as cursor:
        cursor.execute(
            "SELECT id as session_id, title, created_at FROM chat_sessions WHERE user_id=%s ORDER BY created_at DESC",
            (user_id,),
        )
        rows = cursor.fetchall()
        return rows

@app.get("/chat/messages/{session_id}")
def list_chat_messages(session_id: int, conn=Depends(get_db_connection)):
    """Return messages for a session in chronological order."""
    with conn.cursor() as cursor:
        cursor.execute(
            "SELECT role, content, created_at FROM chat_messages WHERE session_id=%s ORDER BY created_at ASC",
            (session_id,),
        )
        rows = cursor.fetchall()
        return rows
@app.post("/auth/google")
def google_auth(google_token: GoogleToken, conn=Depends(get_db_connection)):
    try:
        id_info = id_token.verify_oauth2_token(
            google_token.token, google_requests.Request(), GOOGLE_CLIENT_ID
        )
        user_email = id_info.get("email")
        user_name = id_info.get("name")
        if not user_email:
            raise HTTPException(status_code=400, detail="Email not found in Google token.")

        with conn.cursor() as cursor:
            cursor.execute("SELECT id, name, email FROM next_auth.users WHERE email = %s", (user_email,))
            user = cursor.fetchone()

            if user:
                print(f"Existing user logged in: {user_email}")
            else:
                print(f"Creating new user: {user_email}")
                cursor.execute(
                    "INSERT INTO next_auth.users (name, email) VALUES (%s, %s)",
                    (user_name, user_email)
                )
                conn.commit()

        return {"status": "success", "email": user_email}
    except ValueError:
        raise HTTPException(status_code=401, detail="Invalid Google token.")
    except Exception as e:
        print(f"An unexpected error during Google auth: {e}")
        raise HTTPException(status_code=500, detail="An internal error occurred.")

@app.get("/")
async def root():
    return {"message": "FastAPI backend is running"}

@app.post("/chat/session")
def create_chat_session(user_id: str = Form(...), title: str = Form("New Chatt"), conn=Depends(get_db_connection)):
    with conn.cursor() as cursor:
        cursor.execute("INSERT INTO chat_sessions (user_id, title) VALUES (%s, %s) RETURNING id, created_at", (user_id, title))
        session = cursor.fetchone()
        conn.commit()
        return {"session_id": session["id"], "created_at": session["created_at"]}

@app.get("/chat/sessions/{user_id}")
def get_user_sessions(user_id: str, conn=Depends(get_db_connection)):
    with conn.cursor() as cursor:
        cursor.execute("SELECT id, title, created_at FROM chat_sessions WHERE user_id=%s ORDER BY created_at DESC", (user_id,))
        sessions = cursor.fetchall()
        return [{"session_id": s["id"], "title": s["title"], "created_at": s["created_at"]} for s in sessions]

@app.get("/chat/messages/{session_id}")
def get_chat_messages(session_id: int, conn=Depends(get_db_connection)):
    with conn.cursor() as cursor:
        cursor.execute("SELECT role, content, created_at FROM chat_messages WHERE session_id=%s ORDER BY created_at ASC", (session_id,))
        messages = cursor.fetchall()
        return [{"role": m["role"], "content": m["content"], "created_at": m["created_at"]} for m in messages]

@app.post("/chat/message")
async def send_chat_message(
    prompt: str = Form(...),
    user_id: str = Form(...),
    guest: bool = Form(...),
    session_id: int | None = Form(None),
    file: UploadFile | None = File(None),
    style: str | None = Form(None),
    conn=Depends(get_db_connection)
):
    if file:
        try:
            raw = await file.read()
            file_text = raw.decode("utf-8", errors="ignore")
            prompt += f"\n\n--- Attached file ({file.filename}) ---\n{file_text[:2500]}"
        except Exception as e:
            print("Error reading uploaded file:", e)

    answer = call_llm(prompt, style=style)
    if guest:
        return {"session_id": None, "response": answer}

    with conn.cursor() as cursor:
        if not session_id:
            title = (prompt[:30] + "...") if len(prompt) > 30 else prompt
            if not title.strip():
                title = file.filename[:30] + "..." if file else "New chat"
            cursor.execute("INSERT INTO chat_sessions (user_id, title) VALUES (%s, %s) RETURNING id", (user_id, title))
            session_id = cursor.fetchone()["id"]
        else:
            cursor.execute("SELECT id FROM chat_sessions WHERE id=%s", (session_id,))
            if not cursor.fetchone():
                raise HTTPException(status_code=404, detail="Session not found")

        cursor.execute("INSERT INTO chat_messages (session_id, role, content) VALUES (%s, %s, %s)", (session_id, "user", prompt))
        cursor.execute("INSERT INTO chat_messages (session_id, role, content) VALUES (%s, %s, %s)", (session_id, "bot", answer))
        conn.commit()

    return {"session_id": session_id, "response": answer}

@app.patch("/chat/session/{session_id}")
def update_session_title(session_id: int, title: str = Body(...), conn=Depends(get_db_connection)):
    with conn.cursor() as cursor:
        cursor.execute("UPDATE chat_sessions SET title=%s WHERE id=%s", (title, session_id))
        conn.commit()
        return {"session_id": session_id, "title": title}

@app.get("/chat/search")
def search_user_questions(user_id: str, q: str, limit: int = 20, conn=Depends(get_db_connection)):
    """Search user-asked messages across all sessions for a user.
    Returns session id, session title, matching content snippet and timestamp.
    """
    if not user_id or not q:
        raise HTTPException(status_code=400, detail="user_id and q are required")
    like = f"%{q}%"
    with conn.cursor() as cursor:
        cursor.execute(
            (
                "SELECT s.id AS session_id, s.title, m.content, m.created_at "
                "FROM chat_messages m "
                "JOIN chat_sessions s ON m.session_id = s.id "
                "WHERE s.user_id=%s AND m.role='user' AND m.content ILIKE %s "
                "ORDER BY m.created_at DESC LIMIT %s"
            ),
            (user_id, like, limit),
        )
        rows = cursor.fetchall()
        return [
            {
                "session_id": r["session_id"],
                "title": r["title"],
                "content": r["content"],
                "created_at": r.get("created_at"),
            }
            for r in rows
        ]

@app.delete("/chat/session/{session_id}")
def delete_chat_session(session_id: int, conn=Depends(get_db_connection)):
    """Delete a chat session and its messages."""
    try:
        with conn.cursor() as cursor:
            # Delete dependent messages first to satisfy FK constraints
            cursor.execute("DELETE FROM chat_messages WHERE session_id=%s", (session_id,))
            cursor.execute("DELETE FROM chat_sessions WHERE id=%s", (session_id,))
            conn.commit()
        return {"success": True}
    except Exception as e:
        print("Error deleting chat session:", e)
        raise HTTPException(status_code=500, detail="Failed to delete session")

@app.post("/chat/stream")
async def stream_chat_message(request: Request):
    """Stream assistant response, while persisting user/bot messages for signed-in users.

    Request JSON body supports: { prompt: str, user_id?: str, guest?: bool, session_id?: int }
    If guest is false and session_id is omitted, a new session is created and returned via header X-Session-Id.
    """
    body = await request.json()
    prompt = body.get("prompt")
    user_id = body.get("user_id")
    guest = bool(body.get("guest", False))
    session_id = body.get("session_id")
    style = body.get("style")

    if not prompt:
        raise HTTPException(status_code=400, detail="Prompt is required")

    # For signed-in users, ensure a session exists and persist the user message first
    header_session_id = None
    conn = None
    try:
        if not guest:
            conn = db_pool.getconn()
            with conn.cursor() as cursor:
                if not session_id:
                    title = (prompt[:30] + "...") if len(prompt) > 30 else prompt
                    if not title.strip():
                        title = "New chat"
                    cursor.execute(
                        "INSERT INTO chat_sessions (user_id, title) VALUES (%s, %s) RETURNING id",
                        (user_id, title),
                    )
                    session_id = cursor.fetchone()["id"]
                header_session_id = session_id

                # Persist the user message immediately
                cursor.execute(
                    "INSERT INTO chat_messages (session_id, role, content) VALUES (%s, %s, %s)",
                    (session_id, "user", prompt),
                )
                conn.commit()
    except Exception as e:
        # If persistence setup failed for signed-in user, abort early
        if conn:
            db_pool.putconn(conn)
        raise HTTPException(status_code=500, detail=f"Failed to init chat session: {e}")

    # Wrap the LLM stream to both yield tokens and accumulate full answer
    def wrapper_gen():
        full_answer = ""
        try:
            for chunk in stream_llm_response(prompt, style=style):
                full_answer += chunk
                yield chunk
        finally:
            # On stream completion, persist the assistant message for signed-in users
            if not guest and header_session_id is not None:
                try:
                    with (conn or db_pool.getconn()) as conn_ctx:
                        with conn_ctx.cursor() as cursor:
                            cursor.execute(
                                "INSERT INTO chat_messages (session_id, role, content) VALUES (%s, %s, %s)",
                                (header_session_id, "bot", full_answer),
                            )
                            conn_ctx.commit()
                except Exception as e:
                    # Log and ignore persistence failure after streaming
                    print("Error saving streamed assistant message:", e)
                finally:
                    if conn:
                        db_pool.putconn(conn)

    response = StreamingResponse(wrapper_gen(), media_type="text/event-stream")
    if header_session_id is not None:
        response.headers["X-Session-Id"] = str(header_session_id)
    return response

@app.get("/chat/download/{session_id}")
def download_chat_session(session_id: int, format: str, conn=Depends(get_db_connection)):
    messages = []
    with conn.cursor() as cursor:
        cursor.execute(
            "SELECT role, content FROM chat_messages WHERE session_id=%s ORDER BY created_at ASC",
            (session_id,)
        )
        messages = cursor.fetchall()

    if not messages:
        raise HTTPException(status_code=404, detail="Session not found or has no messages.")

    filename = f"chat_session_{session_id}.{format}"

    if format == "docx":
        file_stream = _create_docx(messages)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif format == "pdf":
        file_stream = _create_pdf(messages)
        media_type = "application/pdf"
    elif format == "csv":
        file_stream = _create_csv(messages)
        media_type = "text/csv"
    else:
        raise HTTPException(status_code=400, detail="Invalid format requested.")

    headers = {'Content-Disposition': f'attachment; filename="{filename}"'}
    return StreamingResponse(file_stream, media_type=media_type, headers=headers)


@app.post("/feedback")
def submit_feedback(body: FeedbackBody, conn=Depends(get_db_connection)):
    """Accept simple user feedback and store in the database.

    Creates the `user_feedback` table if it does not exist.
    """
    msg = (body.message or "").strip()
    if not msg:
        raise HTTPException(status_code=400, detail="Feedback message is required")

    # Normalize rating into 1..5 if provided
    rating = None
    if body.rating is not None:
        try:
            rating = max(1, min(5, int(body.rating)))
        except Exception:
            rating = None

    try:
        with conn.cursor() as cursor:
            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS user_feedback (
                    id SERIAL PRIMARY KEY,
                    user_id TEXT,
                    rating INT,
                    category TEXT,
                    message TEXT NOT NULL,
                    contact_email TEXT,
                    created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
                )
                """
            )
            cursor.execute(
                (
                    "INSERT INTO user_feedback (user_id, rating, category, message, contact_email)\n"
                    "VALUES (%s, %s, %s, %s, %s) RETURNING id, created_at"
                ),
                (body.user_id, rating, (body.category or None), msg, (body.contact_email or None)),
            )
            row = cursor.fetchone()
            conn.commit()
        return {"success": True, "id": row.get("id"), "created_at": row.get("created_at")}
    except Exception as e:
        print("Error saving feedback:", e)
        raise HTTPException(status_code=500, detail="Failed to save feedback")
